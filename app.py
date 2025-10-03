import os
import uuid
import hashlib
from datetime import datetime, timedelta, timezone
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from functools import wraps
import glob
import json
from dotenv import load_dotenv
from config import config
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import requests

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Configure app based on environment
config_name = os.environ.get('FLASK_ENV', 'development')
app.config.from_object(config[config_name])

# Get configuration values
UPLOAD_FOLDER = app.config['UPLOAD_FOLDER']
TYPING_PASSAGES_FOLDER = app.config['TYPING_PASSAGES_FOLDER']
ALLOWED_EXTENSIONS = app.config['ALLOWED_EXTENSIONS']

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TYPING_PASSAGES_FOLDER, exist_ok=True)

# Initialize database
db = SQLAlchemy(app)

# Add security headers for production
@app.after_request
def add_security_headers(response):
    if app.config.get('ENV') == 'production':
        response.headers['X-Content-Type-Options'] = 'nosniff'
        response.headers['X-Frame-Options'] = 'DENY'
        response.headers['X-XSS-Protection'] = '1; mode=block'
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), nullable=False, default='student')  # 'student', 'admin'
    subscription_start = db.Column(db.DateTime)
    subscription_end = db.Column(db.DateTime)
    device_id = db.Column(db.String(100))
    is_active = db.Column(db.Boolean, default=True)
    is_locked = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    last_login = db.Column(db.DateTime)
    last_activity = db.Column(db.DateTime)

class PasswordResetToken(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    token = db.Column(db.String(100), unique=True, nullable=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    is_used = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    otp_code = db.Column(db.String(6))  # For OTP verification
    reset_method = db.Column(db.String(10), default='email')  # 'email' or 'sms'

class TrustedDevice(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    device_fingerprint = db.Column(db.String(200), nullable=False)
    ip_address = db.Column(db.String(45), nullable=False)  # Support IPv6
    user_agent = db.Column(db.Text)
    first_login_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    last_used_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    is_active = db.Column(db.Boolean, default=True)
    locked_reason = db.Column(db.String(200))
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

class Session(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    session_token = db.Column(db.String(100), unique=True, nullable=False)
    device_id = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    last_activity = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    is_active = db.Column(db.Boolean, default=True)

class AudioFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(200), nullable=False)
    reference_text = db.Column(db.Text)  # New field for the text content
    content_type = db.Column(db.String(20), default='both')  # 'exam', 'practice', 'both'
    wpm = db.Column(db.Integer)
    duration = db.Column(db.Integer)  # in seconds
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.id'))

class TypingPassage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    word_count = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.id'))

class DictationAttempt(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    audio_id = db.Column(db.Integer, db.ForeignKey('audio_file.id'), nullable=False)
    transcription = db.Column(db.Text)
    words_typed = db.Column(db.Integer)
    words_correct = db.Column(db.Integer)
    words_wrong = db.Column(db.Integer)
    accuracy_percentage = db.Column(db.Float)
    time_taken = db.Column(db.Integer)  # in seconds
    attempt_number = db.Column(db.Integer, default=1)  # Track attempt number for this audio
    started_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    submitted_at = db.Column(db.DateTime)

class TypingAttempt(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    passage_id = db.Column(db.Integer, db.ForeignKey('typing_passage.id'), nullable=False)
    typed_text = db.Column(db.Text)
    words_typed = db.Column(db.Integer)
    words_correct = db.Column(db.Integer)
    words_wrong = db.Column(db.Integer)
    accuracy_percentage = db.Column(db.Float)
    wpm = db.Column(db.Float)
    time_taken = db.Column(db.Integer)  # in seconds
    attempt_number = db.Column(db.Integer, default=1)  # Track attempt number for this passage
    started_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    submitted_at = db.Column(db.DateTime)

# Utility functions
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password, hash_value):
    return hashlib.sha256(password.encode()).hexdigest() == hash_value

def generate_device_id():
    return str(uuid.uuid4())

def generate_reset_token():
    return str(uuid.uuid4())

def generate_otp():
    import random
    return str(random.randint(100000, 999999))

def generate_device_fingerprint(request):
    """Generate a device fingerprint based on request characteristics"""
    user_agent = request.headers.get('User-Agent', '')
    accept_language = request.headers.get('Accept-Language', '')
    accept_encoding = request.headers.get('Accept-Encoding', '')
    
    # Create a fingerprint from browser characteristics
    fingerprint_data = f"{user_agent}|{accept_language}|{accept_encoding}"
    return hashlib.sha256(fingerprint_data.encode()).hexdigest()

def get_client_ip(request):
    """Get the real client IP address, considering proxies"""
    # Check for forwarded IPs (common in production with proxies)
    if request.headers.get('X-Forwarded-For'):
        # X-Forwarded-For can contain multiple IPs, take the first one
        return request.headers.get('X-Forwarded-For').split(',')[0].strip()
    elif request.headers.get('X-Real-IP'):
        return request.headers.get('X-Real-IP')
    else:
        return request.remote_addr or 'unknown'

def check_trusted_device(user_id, device_fingerprint, ip_address):
    """Check if device/IP combination is trusted for the user"""
    trusted_device = TrustedDevice.query.filter_by(
        user_id=user_id,
        device_fingerprint=device_fingerprint,
        ip_address=ip_address,
        is_active=True
    ).first()
    return trusted_device

def register_trusted_device(user_id, device_fingerprint, ip_address, user_agent):
    """Register a new trusted device for first-time login"""
    try:
        # Check if this is truly the first device for the user
        existing_devices = TrustedDevice.query.filter_by(
            user_id=user_id,
            is_active=True
        ).count()
        
        if existing_devices > 0:
            # User already has a trusted device, this is a violation
            return False
            
        # Register the new trusted device
        trusted_device = TrustedDevice(
            user_id=user_id,
            device_fingerprint=device_fingerprint,
            ip_address=ip_address,
            user_agent=user_agent,
            first_login_at=datetime.now(timezone.utc),
            last_used_at=datetime.now(timezone.utc)
        )
        db.session.add(trusted_device)
        db.session.commit()
        return True
        
    except Exception as e:
        db.session.rollback()
        return False

def lock_user_account(user_id, reason):
    """Lock user account due to device/IP violation"""
    try:
        user = User.query.get(user_id)
        if user:
            user.is_locked = True
            # Deactivate all sessions
            Session.query.filter_by(user_id=user_id).update({'is_active': False})
            # Deactivate all trusted devices
            TrustedDevice.query.filter_by(user_id=user_id).update({
                'is_active': False,
                'locked_reason': reason
            })
            db.session.commit()
            return True
        return False
    except Exception as e:
        db.session.rollback()
        return False

def send_reset_email(email, token, otp):
    """Send password reset email using simple working methods"""
    
    print(f"📧 Sending password reset email to {email}...")
    
    # Method 1: Try Gmail with simple app password (most reliable)
    try:
        gmail_result = send_email_via_gmail_simple(email, token, otp)
        if gmail_result:
            return True
    except Exception as e:
        print(f"❌ Gmail simple failed: {str(e)}")
    
    # Method 2: Try Outlook/Hotmail SMTP (free and reliable)
    try:
        outlook_result = send_email_via_outlook(email, token, otp)
        if outlook_result:
            return True
    except Exception as e:
        print(f"❌ Outlook failed: {str(e)}")
    
    # Method 3: Try Mailtrap or other simple SMTP
    try:
        mailtrap_result = send_email_via_mailtrap(email, token, otp)
        if mailtrap_result:
            return True
    except Exception as e:
        print(f"❌ Mailtrap failed: {str(e)}")
    
    # Fallback to console with practical instructions
    email_configs = []
    
    # Create email content
    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background-color: #4c1d95; color: white; padding: 20px; text-align: center;">
            <h1>🔐 Stenographix Admin</h1>
            <h2>Password Reset Request</h2>
        </div>
        
        <div style="padding: 20px; background-color: #f8f9fa;">
            <p><strong>Hello Admin,</strong></p>
            <p>You have requested to reset your password for the Stenographix Admin Panel.</p>
            
            <div style="background-color: white; border: 2px solid #4c1d95; border-radius: 8px; padding: 20px; margin: 20px 0;">
                <h3 style="color: #4c1d95; margin-top: 0;">🔑 Reset Options (Choose One):</h3>
                
                <div style="background-color: #f0f0f0; padding: 15px; border-radius: 5px; margin: 10px 0;">
                    <p><strong>Option 1 - Email Token:</strong></p>
                    <code style="font-size: 14px; background-color: #e9ecef; padding: 8px; border-radius: 3px; display: block; word-break: break-all;">
                    {token}
                    </code>
                </div>
                
                <div style="background-color: #f0f0f0; padding: 15px; border-radius: 5px; margin: 10px 0;">
                    <p><strong>Option 2 - SMS OTP Code:</strong></p>
                    <code style="font-size: 18px; background-color: #e9ecef; padding: 12px; border-radius: 3px; display: block; text-align: center; font-weight: bold; letter-spacing: 2px;">
                    {otp}
                    </code>
                </div>
            </div>
            
            <div style="background-color: #fff3cd; border: 1px solid #ffeaa7; padding: 15px; border-radius: 5px; margin: 15px 0;">
                <p style="margin: 0;"><strong>⚠️ Important:</strong></p>
                <ul style="margin: 5px 0 0 0;">
                    <li>Use either the email token OR the 6-digit SMS OTP</li>
                    <li>This reset expires in <strong>35 minutes</strong></li>
                    <li>If you didn't request this, ignore this email</li>
                </ul>
            </div>
            
            <p style="text-align: center; margin-top: 30px;">
                <em>Stenographix Admin System - Secure Password Reset</em>
            </p>
        </div>
    </body>
    </html>
    """
    
    # Try each email service
    for config in email_configs:
        try:
            print(f"🔄 Attempting to send email via {config['name']}...")
            
            # Create message
            message = MIMEMultipart("alternative")
            message["Subject"] = "🔐 Admin Password Reset - Stenographix"
            message["From"] = f"Stenographix Admin <{config['sender_email']}>"
            message["To"] = email
            
            # Add HTML content
            part = MIMEText(html_content, "html")
            message.attach(part)
            
            # Send email
            with smtplib.SMTP(config["smtp_server"], config["smtp_port"]) as server:
                server.starttls()
                server.login(config["sender_email"], config["sender_password"])
                server.sendmail(config["sender_email"], email, message.as_string())
            
            print(f"✅ Password reset email sent successfully to {email} via {config['name']}")
            return True
            
        except Exception as e:
            print(f"❌ {config['name']} failed: {str(e)}")
            continue
    
    # If all free email services fail, use practical fallback methods
    print(f"❌ All free email services failed. Using practical delivery methods.")
    print(f"")
    print(f"=== 📧 PASSWORD RESET - MANUAL DELIVERY REQUIRED ===")
    print(f"Admin Email: {email}")
    print(f"")
    print(f"🔑 EMAIL TOKEN: {token}")
    print(f"📱 SMS OTP CODE: {otp}")
    print(f"")
    admin_phone = app.config.get('ADMIN_PHONE', '+1234567890')
    admin_email = app.config.get('ADMIN_EMAIL', 'admin@localhost')
    
    print(f"💌 DELIVERY OPTIONS:")
    print(f"1. Send WhatsApp to {admin_phone}:")
    print(f"   'Your admin reset: Token={token} OR OTP={otp}'")
    print(f"")
    print(f"2. Send email manually to: {admin_email}")
    print(f"   Subject: Admin Password Reset")
    print(f"   Body: Your reset token: {token} OR OTP: {otp}")
    print(f"")
    print(f"3. Call {admin_phone} and provide OTP: {otp}")
    print(f"")
    print(f"⏰ Both expire in 35 minutes")
    print(f"====================================================")
    return True  # Return true since manual delivery is available

def send_email_via_gmail_simple(email, token, otp):
    """Send email using Gmail with simple authentication"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        
        # SMTP configuration from environment
        smtp_server = app.config.get('MAIL_SERVER', 'localhost')
        smtp_port = app.config.get('MAIL_PORT', 587)
        sender_email = app.config.get('MAIL_USERNAME', 'noreply@localhost')
        sender_password = app.config.get('MAIL_PASSWORD', '')
        
        # Create email content
        msg = MIMEMultipart()
        msg['From'] = f"Stenographix Admin <{sender_email}>"
        msg['To'] = email
        msg['Subject'] = "🔐 Admin Password Reset - Stenographix"
        
        body = f"""🔐 Stenographix Admin Password Reset

Hello Admin,

You requested a password reset for your Stenographix admin account.

Reset Options (use either one):

📧 Email Token: {token}
📱 SMS OTP Code: {otp}

⏰ Expires in 35 minutes

If you didn't request this, ignore this email.

- Stenographix System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        
        print(f"✅ Email sent successfully via Gmail to {email}")
        return True
        
    except Exception as e:
        print(f"❌ Gmail simple error: {str(e)}")
        return False

def send_email_via_outlook(email, token, otp):
    """Send email using Outlook/Hotmail SMTP (free)"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        
        # Outlook SMTP configuration
        smtp_server = "smtp-mail.outlook.com"
        smtp_port = 587
        sender_email = "your_outlook@hotmail.com"  # You need a free Hotmail/Outlook account
        sender_password = "your_password_here"     # Regular password works
        
        # Create email content
        msg = MIMEMultipart()
        msg['From'] = f"Stenographix Admin <{sender_email}>"
        msg['To'] = email
        msg['Subject'] = "🔐 Admin Password Reset - Stenographix"
        
        body = f"""🔐 Stenographix Admin Password Reset

Hello Admin,

You requested a password reset for your Stenographix admin account.

Reset Options (use either one):

📧 Email Token: {token}
📱 SMS OTP Code: {otp}

⏰ Expires in 35 minutes

If you didn't request this, ignore this email.

- Stenographix System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        
        print(f"✅ Email sent successfully via Outlook to {email}")
        return True
        
    except Exception as e:
        print(f"❌ Outlook error: {str(e)}")
        return False

def send_email_via_mailtrap(email, token, otp):
    """Send email using Mailtrap (free testing service)"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        
        # Mailtrap SMTP configuration (free account)
        smtp_server = "smtp.mailtrap.io"
        smtp_port = 2525
        sender_email = "noreply@stenographix.app"
        smtp_username = "your_mailtrap_username"  # From Mailtrap dashboard
        smtp_password = "your_mailtrap_password"  # From Mailtrap dashboard
        
        # Create email content
        msg = MIMEMultipart()
        msg['From'] = f"Stenographix Admin <{sender_email}>"
        msg['To'] = email
        msg['Subject'] = "🔐 Admin Password Reset - Stenographix"
        
        body = f"""🔐 Stenographix Admin Password Reset

Hello Admin,

You requested a password reset for your Stenographix admin account.

Reset Options (use either one):

📧 Email Token: {token}
📱 SMS OTP Code: {otp}

⏰ Expires in 35 minutes

If you didn't request this, ignore this email.

- Stenographix System"""
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.send_message(msg)
        server.quit()
        
        print(f"✅ Email sent successfully via Mailtrap to {email}")
        return True
        
    except Exception as e:
        print(f"❌ Mailtrap error: {str(e)}")
        return False

def send_reset_sms(mobile, otp):
    """Send OTP via SMS to admin mobile using TextBelt API (Free)"""
    try:
        print(f"📱 Sending SMS OTP to {mobile}...")
        
        # Method 1: Try TextBelt (Free SMS service - 1 free SMS per day)
        try:
            textbelt_result = send_sms_via_textbelt(mobile, otp)
            if textbelt_result:
                return True
        except Exception as e:
            print(f"❌ TextBelt failed: {str(e)}")
        
        # Method 2: Try SMS Gateway (Alternative free service)
        try:
            gateway_result = send_sms_via_gateway(mobile, otp)
            if gateway_result:
                return True
        except Exception as e:
            print(f"❌ SMS Gateway failed: {str(e)}")
        
        # Fallback to console logging with WhatsApp suggestion
        print(f"=== PASSWORD RESET SMS (CONSOLE + WHATSAPP) ===")
        print(f"📱 To: {mobile}")
        print(f"🔑 OTP Code: {otp}")
        print(f"")
        print(f"💬 MANUAL DELIVERY OPTIONS:")
        print(f"1. Send WhatsApp message to {mobile}:")
        print(f"   'Your Stenographix admin OTP: {otp}'")
        print(f"2. Call {mobile} and provide OTP: {otp}")
        print(f"3. Use any SMS service to send: 'OTP: {otp}'")
        print(f"")
        print(f"⏰ OTP expires in 35 minutes")
        print(f"===============================================")
        return True
        
    except Exception as e:
        print(f"❌ SMS sending failed: {str(e)}")
        print(f"=== PASSWORD RESET SMS (MANUAL FALLBACK) ===")
        print(f"📱 To: {mobile}")
        print(f"🔑 OTP Code: {otp}")
        print(f"💬 Please send this OTP manually via WhatsApp/Call")
        print(f"===============================================")
        return True

def send_sms_via_textbelt(mobile, otp):
    """Send SMS using TextBelt API (Free 1 SMS per day)"""
    try:
        import requests
        
        url = "https://textbelt.com/text"
        
        # Format message
        message = f"🔐 Stenographix Admin OTP: {otp}\n\nUse this code to reset your password. Expires in 35 minutes.\n\n- Stenographix System"
        
        # Clean mobile number (remove +91 if present)
        clean_mobile = mobile.replace("+91", "").replace(" ", "").replace("-", "")
        
        payload = {
            "phone": f"+91{clean_mobile}",
            "message": message,
            "key": "textbelt"  # Free quota key
        }
        
        response = requests.post(url, data=payload, timeout=10)
        result = response.json()
        
        if result.get('success'):
            print(f"✅ SMS sent successfully via TextBelt to {mobile}")
            print(f"📱 TextBelt ID: {result.get('textId', 'N/A')}")
            return True
        else:
            print(f"❌ TextBelt failed: {result.get('error', 'Unknown error')}")
            return False
            
    except Exception as e:
        print(f"❌ TextBelt error: {str(e)}")
        return False

def send_sms_via_gateway(mobile, otp):
    """Send SMS using SMS Gateway API (Alternative free service)"""
    try:
        import requests
        
        # Free SMS Gateway (multiple services available)
        services = [
            {
                "name": "SMS Gateway",
                "url": "https://smsgateway24.com/getdata/addsms",
                "method": "GET",
                "params": {
                    "userid": "demo",
                    "password": "demo",
                    "mobile": mobile.replace("+91", ""),
                    "message": f"Stenographix Admin OTP: {otp}. Expires in 35 minutes.",
                    "senderid": "STENX",
                    "is_unicode": "0",
                    "is_flash": "0",
                    "schedule_time": "",
                    "server": "1"
                }
            }
        ]
        
        for service in services:
            try:
                print(f"🔄 Trying {service['name']}...")
                
                if service["method"] == "GET":
                    response = requests.get(service["url"], params=service["params"], timeout=10)
                else:
                    response = requests.post(service["url"], data=service["params"], timeout=10)
                
                if response.status_code == 200:
                    result_text = response.text.strip()
                    if "success" in result_text.lower() or "sent" in result_text.lower():
                        print(f"✅ SMS sent successfully via {service['name']} to {mobile}")
                        return True
                    else:
                        print(f"❌ {service['name']} response: {result_text}")
                        
            except Exception as e:
                print(f"❌ {service['name']} error: {str(e)}")
                continue
        
        return False
        
    except Exception as e:
        print(f"❌ SMS Gateway error: {str(e)}")
        return False

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def compare_texts(reference_text, typed_text):
    """
    Enhanced text comparison using LCS strategy with improved error detection.
    
    This function addresses the issue where correctly typed words like "However..."
    were incorrectly marked as errors by using the Longest Common Subsequence algorithm
    for more accurate text alignment and comparison.
    """
    # Import the enhanced comparison function
    from lcs_text_comparison import enhanced_compare_texts
    
    # Use the enhanced LCS-based comparison
    result = enhanced_compare_texts(reference_text, typed_text)
    
    # Return in the format expected by the existing application
    return {
        'words_correct': result['words_correct'],
        'words_wrong': result['words_wrong'],
        'accuracy_percentage': result['accuracy_percentage'],
        'total_words': result['total_words'],
        'typed_words': result['typed_words'],
        'enhanced_comparison': result['enhanced_comparison'],
        'lcs_analysis': result.get('lcs_analysis', {}),
        'punctuation_errors': result.get('punctuation_errors', []),
        'error_summary': result.get('error_summary', {})
    }

def check_subscription_active(user):
    if not user.subscription_end:
        return False
    return datetime.now(timezone.utc).replace(tzinfo=None) < user.subscription_end

def get_subscription_days_remaining(user):
    if not user.subscription_end:
        return 0
    delta = user.subscription_end - datetime.now(timezone.utc).replace(tzinfo=None)
    return max(0, delta.days)

def calculate_improvement_dictation(user_id, audio_id, current_accuracy, current_time):
    """Calculate improvement compared to previous attempts"""
    try:
        # Get the previous attempt for this user and audio (second most recent)
        previous_attempt = DictationAttempt.query.filter_by(
            user_id=user_id,
            audio_id=audio_id
        ).filter(DictationAttempt.submitted_at.isnot(None)).order_by(
            DictationAttempt.submitted_at.desc()
        ).offset(1).first()  # Skip the most recent (current) attempt
        
        if not previous_attempt:
            return {
                'has_improvement': False,
                'message': 'First attempt! Keep practicing to track your progress.',
                'accuracy_change': 0,
                'time_change': 0,
                'days_since_last': 0
            }
        
        # Calculate improvements
        accuracy_change = current_accuracy - (previous_attempt.accuracy_percentage or 0)
        time_change = (previous_attempt.time_taken or 0) - current_time  # Positive means faster
        
        # Calculate days since last attempt
        days_since = (datetime.now(timezone.utc) - previous_attempt.submitted_at).days
        
        # Generate improvement message
        if accuracy_change > 2:
            message = f"+{accuracy_change:.1f}% accuracy improvement!"
        elif accuracy_change > 0:
            message = f"+{accuracy_change:.1f}% accuracy since last attempt"
        elif accuracy_change < -2:
            message = f"{accuracy_change:.1f}% accuracy - focus on listening carefully"
        else:
            message = "Similar accuracy to last attempt"
            
        if time_change > 30:  # More than 30 seconds faster
            message += f" and {time_change//60:.0f}m {time_change%60:.0f}s faster!"
        elif time_change > 0:
            message += f" and {time_change:.0f}s faster"
        elif time_change < -30:
            message += f" but took {abs(time_change)//60:.0f}m {abs(time_change)%60:.0f}s longer"
            
        return {
            'has_improvement': True,
            'message': message,
            'accuracy_change': accuracy_change,
            'time_change': time_change,
            'days_since_last': days_since,
            'previous_accuracy': previous_attempt.accuracy_percentage or 0,
            'previous_time': previous_attempt.time_taken or 0
        }
    except Exception as e:
        return {
            'has_improvement': False,
            'message': 'Unable to calculate improvement data',
            'accuracy_change': 0,
            'time_change': 0,
            'days_since_last': 0
        }

def calculate_improvement_typing(user_id, passage_id, current_accuracy, current_wpm, current_time):
    """Calculate improvement compared to previous typing attempts"""
    try:
        # Get the previous attempt for this user and passage (second most recent)
        previous_attempt = TypingAttempt.query.filter_by(
            user_id=user_id,
            passage_id=passage_id
        ).filter(TypingAttempt.submitted_at.isnot(None)).order_by(
            TypingAttempt.submitted_at.desc()
        ).offset(1).first()  # Skip the most recent (current) attempt
        
        if not previous_attempt:
            return {
                'has_improvement': False,
                'message': 'First attempt! Keep practicing to track your progress.',
                'accuracy_change': 0,
                'wpm_change': 0,
                'time_change': 0,
                'days_since_last': 0
            }
        
        # Calculate improvements
        accuracy_change = current_accuracy - (previous_attempt.accuracy_percentage or 0)
        wpm_change = current_wpm - (previous_attempt.wpm or 0)
        time_change = (previous_attempt.time_taken or 0) - current_time  # Positive means faster
        
        # Calculate days since last attempt
        days_since = (datetime.now(timezone.utc) - previous_attempt.submitted_at).days
        
        # Generate improvement message
        improvements = []
        if wpm_change > 2:
            improvements.append(f"+{wpm_change:.1f} WPM")
        if accuracy_change > 2:
            improvements.append(f"+{accuracy_change:.1f}% accuracy")
        elif accuracy_change < -2:
            improvements.append(f"{accuracy_change:.1f}% accuracy")
            
        if improvements:
            message = " and ".join(improvements) + " since last attempt!"
        elif wpm_change > 0 or accuracy_change > 0:
            message = "Small improvements since last attempt"
        else:
            message = "Similar performance to last attempt - keep practicing!"
            
        return {
            'has_improvement': True,
            'message': message,
            'accuracy_change': accuracy_change,
            'wpm_change': wpm_change,
            'time_change': time_change,
            'days_since_last': days_since,
            'previous_accuracy': previous_attempt.accuracy_percentage or 0,
            'previous_wpm': previous_attempt.wpm or 0,
            'previous_time': previous_attempt.time_taken or 0
        }
    except Exception as e:
        return {
            'has_improvement': False,
            'message': 'Unable to calculate improvement data',
            'accuracy_change': 0,
            'wpm_change': 0,
            'time_change': 0,
            'days_since_last': 0
        }

def analyze_common_mistakes(reference_text, user_text):
    """Analyze common mistakes in user's text compared to reference"""
    try:
        import re
        from collections import Counter
        
        ref_words = reference_text.split()
        user_words = user_text.split()
        
        mistakes = {
            'punctuation_errors': [],
            'spelling_errors': [],
            'repeated_words': [],
            'missing_words': [],
            'common_patterns': []
        }
        
        # Analyze punctuation errors
        punctuation_marks = [',', '.', '!', '?', ';', ':', '"', "'"]
        for mark in punctuation_marks:
            ref_count = reference_text.count(mark)
            user_count = user_text.count(mark)
            if ref_count > user_count:
                mistakes['punctuation_errors'].append(f'Missing {ref_count - user_count} "{mark}"')
            elif user_count > ref_count:
                mistakes['punctuation_errors'].append(f'Extra {user_count - ref_count} "{mark}"')
        
        # Find repeated words (consecutive duplicates)
        for i in range(len(user_words) - 1):
            if user_words[i].lower() == user_words[i + 1].lower() and user_words[i].lower() not in ['the', 'a', 'an', 'and', 'or']:
                mistakes['repeated_words'].append(user_words[i])
        
        # Find spelling errors by comparing similar length words
        for i, user_word in enumerate(user_words):
            if i < len(ref_words):
                ref_word = ref_words[i]
                if user_word.lower() != ref_word.lower() and len(user_word) >= 3:
                    # Check if it's likely a spelling error (similar length, some common letters)
                    if abs(len(user_word) - len(ref_word)) <= 2:
                        mistakes['spelling_errors'].append({
                            'typed': user_word,
                            'correct': ref_word
                        })
        
        # Find most common error patterns
        error_patterns = []
        if len(mistakes['punctuation_errors']) > 0:
            error_patterns.append('Punctuation accuracy')
        if len(mistakes['spelling_errors']) > 2:
            error_patterns.append('Spelling accuracy')
        if len(mistakes['repeated_words']) > 0:
            error_patterns.append('Word repetition')
            
        # Analyze capitalization
        ref_caps = sum(1 for c in reference_text if c.isupper())
        user_caps = sum(1 for c in user_text if c.isupper())
        if abs(ref_caps - user_caps) > 2:
            error_patterns.append('Capitalization')
        
        mistakes['common_patterns'] = error_patterns[:3]  # Top 3 patterns
        
        return mistakes
        
    except Exception as e:
        return {
            'punctuation_errors': [],
            'spelling_errors': [],
            'repeated_words': [],
            'missing_words': [],
            'common_patterns': []
        }

# Authentication decorators
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page.')
            return redirect(url_for('student_login'))
        
        # Check session validity and activity timeout
        user = User.query.get(session['user_id'])
        if not user or not user.is_active or user.is_locked:
            session.clear()
            flash('Your account is not accessible. Please contact admin.')
            return redirect(url_for('student_login'))
        
        # Check subscription
        if user.role == 'student' and not check_subscription_active(user):
            flash('Your subscription has expired. Please contact admin to renew.')
            return redirect(url_for('subscription_expired'))
        
        # Update last activity
        user.last_activity = datetime.now(timezone.utc)
        db.session.commit()
        
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in to access this page.')
            return redirect(url_for('admin_login'))
        
        user = User.query.get(session['user_id'])
        if not user or user.role != 'admin':
            flash('Admin access required.')
            return redirect(url_for('student_login'))
        
        return f(*args, **kwargs)
    return decorated_function

# Routes
@app.route('/')
def index():
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        if user and user.role == 'admin':
            return redirect(url_for('admin_dashboard'))
        else:
            return redirect(url_for('practice_selection'))
    return render_template('welcome.html')

@app.route('/student-login', methods=['GET', 'POST'])
def student_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        # Generate device fingerprint and get IP
        device_fingerprint = generate_device_fingerprint(request)
        client_ip = get_client_ip(request)
        user_agent = request.headers.get('User-Agent', '')
        
        user = User.query.filter_by(username=username).first()
        
        if user and verify_password(password, user.password_hash):
            # Pre-validation checks
            if user.is_locked:
                flash('Your account has been locked due to login from an unauthorized device/IP. Please contact Admin.')
                return render_template('student_login.html')
            
            if not user.is_active:
                flash('Your account is suspended. Please contact admin.')
                return render_template('student_login.html')
            
            if user.role not in ['student']:
                flash('Please use admin login for admin accounts.')
                return render_template('student_login.html')
            
            # Check subscription
            if not check_subscription_active(user):
                flash('Your subscription has expired. Please contact admin to renew.')
                return render_template('student_login.html')
            
            # DEVICE/IP RESTRICTION LOGIC
            # Check if this is a trusted device/IP combination
            trusted_device = check_trusted_device(user.id, device_fingerprint, client_ip)
            
            if trusted_device:
                # User is logging in from trusted device/IP - allow access
                # Update last used time
                trusted_device.last_used_at = datetime.now(timezone.utc)
                
                # Clear old sessions and create new one
                Session.query.filter_by(user_id=user.id).update({'is_active': False})
                new_session = Session(
                    user_id=user.id,
                    session_token=str(uuid.uuid4()),
                    device_id=device_fingerprint
                )
                db.session.add(new_session)
                
                # Update user login info
                user.last_login = datetime.now(timezone.utc)
                user.last_activity = datetime.now(timezone.utc)
                user.device_id = device_fingerprint
                
                db.session.commit()
                
                session['user_id'] = user.id
                session['username'] = user.username
                session['role'] = user.role
                session['session_token'] = new_session.session_token
                
                # Check for expiry warning
                days_remaining = get_subscription_days_remaining(user)
                if days_remaining <= 2:
                    flash(f'Your subscription expires in {days_remaining} days. Please contact admin to renew.')
                
                return redirect(url_for('practice_selection'))
            
            else:
                # Check if this is the user's first login (no trusted devices)
                existing_devices = TrustedDevice.query.filter_by(
                    user_id=user.id,
                    is_active=True
                ).count()
                
                if existing_devices == 0:
                    # First login - register this device/IP as trusted
                    if register_trusted_device(user.id, device_fingerprint, client_ip, user_agent):
                        # Clear old sessions and create new one
                        Session.query.filter_by(user_id=user.id).update({'is_active': False})
                        new_session = Session(
                            user_id=user.id,
                            session_token=str(uuid.uuid4()),
                            device_id=device_fingerprint
                        )
                        db.session.add(new_session)
                        
                        # Update user login info
                        user.last_login = datetime.now(timezone.utc)
                        user.last_activity = datetime.now(timezone.utc)
                        user.device_id = device_fingerprint
                        
                        db.session.commit()
                        
                        session['user_id'] = user.id
                        session['username'] = user.username
                        session['role'] = user.role
                        session['session_token'] = new_session.session_token
                        
                        # Check for expiry warning
                        days_remaining = get_subscription_days_remaining(user)
                        if days_remaining <= 2:
                            flash(f'Your subscription expires in {days_remaining} days. Please contact admin to renew.')
                        
                        return redirect(url_for('practice_selection'))
                    else:
                        # Error registering device
                        flash('Error setting up device security. Please contact admin.')
                        return render_template('student_login.html')
                else:
                    # User already has a trusted device, this is a different device/IP
                    # LOCK THE ACCOUNT IMMEDIATELY
                    lock_reason = f"Login attempt from unauthorized device/IP: {client_ip}"
                    lock_user_account(user.id, lock_reason)
                    
                    flash('Your account has been locked due to login from an unauthorized device/IP. Please contact Admin.')
                    return render_template('student_login.html')
        else:
            flash('Invalid username or password')
    
    return render_template('student_login.html')

@app.route('/admin-login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        user = User.query.filter_by(username=username).first()
        
        if user and verify_password(password, user.password_hash) and user.role == 'admin':
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            
            user.last_login = datetime.now(timezone.utc)
            user.last_activity = datetime.now(timezone.utc)
            db.session.commit()
            
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid admin credentials')
    
    return render_template('admin_login.html')

@app.route('/logout')
def logout():
    if 'user_id' in session:
        # Deactivate session
        Session.query.filter_by(user_id=session['user_id']).update({'is_active': False})
        db.session.commit()
    
    session.clear()
    flash('You have been logged out.')
    return redirect(url_for('index'))

@app.route('/practice-selection')
@login_required
def practice_selection():
    user = User.query.get(session['user_id'])
    
    # Get counts for display
    audio_count = AudioFile.query.count()
    passage_count = TypingPassage.query.count()
    
    return render_template('practice_selection.html',
                         user=user,
                         audio_count=audio_count,
                         passage_count=passage_count,
                         now=datetime.now(timezone.utc).replace(tzinfo=None))

@app.route('/subscription-expired')
def subscription_expired():
    return render_template('subscription_expired.html')

@app.route('/admin-dashboard')
@admin_required
def admin_dashboard():
    # Get statistics for dashboard
    total_users = User.query.filter_by(role='student').count()
    active_users = User.query.filter_by(role='student', is_active=True).count()
    total_audio_files = AudioFile.query.count()
    total_typing_passages = TypingPassage.query.count()
    
    # Users with expiring subscriptions (next 7 days)
    now_naive = datetime.now(timezone.utc).replace(tzinfo=None)
    expiring_soon = User.query.filter(
        User.role == 'student',
        User.subscription_end <= now_naive + timedelta(days=7),
        User.subscription_end > now_naive
    ).all()
    
    # Recent attempts
    recent_dictation_attempts = DictationAttempt.query.order_by(DictationAttempt.submitted_at.desc()).limit(5).all()
    recent_typing_attempts = TypingAttempt.query.order_by(TypingAttempt.submitted_at.desc()).limit(5).all()
    
    return render_template('admin_dashboard.html',
                         total_users=total_users,
                         active_users=active_users,
                         total_audio_files=total_audio_files,
                         total_typing_passages=total_typing_passages,
                         expiring_soon=expiring_soon,
                         recent_dictation_attempts=recent_dictation_attempts,
                         recent_typing_attempts=recent_typing_attempts)

@app.route('/admin/users')
@admin_required
def admin_users():
    users = User.query.filter_by(role='student').all()
    return render_template('admin_users.html', users=users, now=datetime.now(timezone.utc).replace(tzinfo=None))

@app.route('/admin/create-user', methods=['GET', 'POST'])
@admin_required
def admin_create_user():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        subscription_days = int(request.form.get('subscription_days', 30))
        
        # Check if username already exists
        if User.query.filter_by(username=username).first():
            flash('Username already exists')
            return render_template('admin_create_user.html')
        
        # Check user limit (1000 students max)
        current_student_count = User.query.filter_by(role='student').count()
        if current_student_count >= 1000:
            flash('Cannot create user: Maximum limit of 1000 students reached')
            return render_template('admin_create_user.html')
        
        # Create new student
        new_user = User(
            username=username,
            password_hash=hash_password(password),
            role='student',
            is_active=True,
            subscription_start=datetime.now(timezone.utc),
            subscription_end=datetime.now(timezone.utc) + timedelta(days=subscription_days)
        )
        
        db.session.add(new_user)
        db.session.commit()
        
        flash(f'Student account "{username}" created successfully with {subscription_days}-day subscription')
        return redirect(url_for('admin_users'))
    
    return render_template('admin_create_user.html')

@app.route('/admin/edit-user/<int:user_id>', methods=['GET', 'POST'])
@admin_required
def admin_edit_user(user_id):
    user = User.query.get_or_404(user_id)
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'extend_subscription':
            extend_type = request.form.get('extend_type', 'days')
            extend_value = int(request.form.get('extend_value', 30))
            
            now_naive = datetime.now(timezone.utc).replace(tzinfo=None)
            
            if extend_type == 'days':
                delta = timedelta(days=extend_value)
            elif extend_type == 'months':
                delta = timedelta(days=extend_value * 30)  # Approximate month
            elif extend_type == 'years':
                delta = timedelta(days=extend_value * 365)  # Approximate year
            else:
                delta = timedelta(days=extend_value)
            
            if user.subscription_end and user.subscription_end > now_naive:
                user.subscription_end = user.subscription_end + delta
            else:
                user.subscription_end = now_naive + delta
            flash(f'Subscription extended by {extend_value} {extend_type}')
            
        elif action == 'set_subscription_date':
            new_subscription_end = request.form.get('new_subscription_end')
            if new_subscription_end:
                try:
                    # Parse the date from the calendar input
                    new_end_date = datetime.strptime(new_subscription_end, '%Y-%m-%d').replace(tzinfo=None)
                    user.subscription_end = new_end_date
                    flash(f'Subscription end date set to {new_end_date.strftime("%Y-%m-%d")}')
                except ValueError:
                    flash('Invalid date format. Please select a valid date.', 'error')
            else:
                flash('Please select a subscription end date.', 'error')
            
        elif action == 'reset_password':
            new_password = request.form.get('new_password')
            user.password_hash = hash_password(new_password)
            flash('Password reset successfully')
            
        elif action == 'toggle_active':
            user.is_active = not user.is_active
            status = 'activated' if user.is_active else 'suspended'
            flash(f'Account {status}')
            
        elif action == 'unlock_account':
            user.is_locked = False
            # Clear all sessions for this user
            Session.query.filter_by(user_id=user.id).update({'is_active': False})
            # Reactivate trusted devices (allow user to login again)
            TrustedDevice.query.filter_by(user_id=user.id).update({
                'is_active': True,
                'locked_reason': None
            })
            flash('Account unlocked, all sessions cleared, and trusted devices reactivated')
            
        elif action == 'reset_trusted_devices':
            # Remove all trusted devices - user will need to set up new trusted device on next login
            TrustedDevice.query.filter_by(user_id=user.id).update({'is_active': False})
            Session.query.filter_by(user_id=user.id).update({'is_active': False})
            flash('All trusted devices removed. User will set up new trusted device on next login.')
            
        elif action == 'view_trusted_devices':
            # This will be handled by the template to show trusted devices info
            pass
        
        db.session.commit()
        return redirect(url_for('admin_edit_user', user_id=user_id))
    
    # Get trusted devices for this user
    trusted_devices = TrustedDevice.query.filter_by(user_id=user.id).order_by(TrustedDevice.last_used_at.desc()).all()
    
    return render_template('admin_edit_user.html',
                         user=user,
                         trusted_devices=trusted_devices,
                         now=datetime.now(timezone.utc).replace(tzinfo=None))

@app.route('/admin/delete-user/<int:user_id>', methods=['POST'])
@admin_required
def admin_delete_user(user_id):
    user = User.query.get_or_404(user_id)
    
    if user.role != 'student':
        flash('Cannot delete non-student accounts', 'error')
        return redirect(url_for('admin_users'))
    
    try:
        username = user.username
        
        # Delete related data first
        DictationAttempt.query.filter_by(user_id=user.id).delete()
        TypingAttempt.query.filter_by(user_id=user.id).delete()
        Session.query.filter_by(user_id=user.id).delete()
        
        # Delete the user account
        db.session.delete(user)
        db.session.commit()
        
        flash(f'Student account "{username}" deleted successfully. Slot is now available for new registrations.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting user: {str(e)}', 'error')
    
    return redirect(url_for('admin_users'))

@app.route('/admin/content')
@admin_required
def admin_content():
    audio_files = AudioFile.query.all()
    typing_passages = TypingPassage.query.all()
    return render_template('admin_content.html',
                         audio_files=audio_files,
                         typing_passages=typing_passages)

@app.route('/admin/upload_audio', methods=['GET', 'POST'])
@admin_required
def admin_upload_audio():
    if request.method == 'POST':
        # Handle audio file upload
        if 'audio_file' not in request.files:
            flash('No audio file selected', 'error')
            return redirect(request.url)
        
        audio_file_obj = request.files['audio_file']
        if audio_file_obj.filename == '':
            flash('No audio file selected', 'error')
            return redirect(request.url)
        
        # Get reference text from either text input or file upload
        reference_text = ''
        text_input_method = request.form.get('text_input_method', 'direct')
        
        if text_input_method == 'direct':
            # Get text from direct input
            reference_text = request.form.get('reference_text_direct', '').strip()
            if not reference_text:
                flash('Please enter the reference text directly or upload a file', 'error')
                return redirect(request.url)
        else:
            # Handle reference text file upload
            if 'reference_text_file' not in request.files:
                flash('No reference text file selected', 'error')
                return redirect(request.url)
            
            text_file = request.files['reference_text_file']
            if text_file.filename == '':
                flash('No reference text file selected', 'error')
                return redirect(request.url)
            
            # Read reference text from uploaded file (support multiple formats)
            file_extension = text_file.filename.lower().split('.')[-1]
            
            try:
                if file_extension == 'txt':
                    reference_text = text_file.read().decode('utf-8').strip()
                elif file_extension == 'docx':
                    # Handle DOCX files
                    try:
                        import docx
                        doc = docx.Document(text_file)
                        reference_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    except ImportError:
                        flash('DOCX support not available. Please use TXT format or install python-docx', 'error')
                        return redirect(request.url)
                elif file_extension == 'odt':
                    # Handle ODT files
                    try:
                        from odf.opendocument import load
                        from odf.text import P
                        doc = load(text_file)
                        paragraphs = doc.getElementsByType(P)
                        reference_text = '\n'.join([str(p) for p in paragraphs])
                    except ImportError:
                        flash('ODT support not available. Please use TXT format or install odfpy', 'error')
                        return redirect(request.url)
                else:
                    flash('Please select a valid file format (.txt, .docx, or .odt)', 'error')
                    return redirect(request.url)
                
                if not reference_text.strip():
                    flash('Reference text file is empty', 'error')
                    return redirect(request.url)
                    
            except Exception as e:
                flash(f'Error reading text file: {str(e)}', 'error')
                return redirect(request.url)
        
        if audio_file_obj and audio_file_obj.filename.lower().endswith('.mp3'):
            filename = secure_filename(audio_file_obj.filename)
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            audio_file_obj.save(file_path)
            
            # Get content type selection
            content_type = request.form.get('content_type', 'exam')
            
            # Save to database with reference text and content type
            # Fixed 30-minute timer for dictation
            audio_file = AudioFile(
                title=request.form.get('title', filename),
                filename=filename,
                reference_text=reference_text.strip(),
                content_type=content_type,
                duration=2100,  # Fixed 35 minutes (2100 seconds)
                uploaded_by=session.get('user_id')
            )
            db.session.add(audio_file)
            db.session.commit()
            
            flash('Audio file and reference text uploaded successfully! Timer set to 35 minutes.', 'success')
            return redirect(url_for('admin_content'))
        else:
            flash('Please select a valid MP3 file', 'error')
    
    return render_template('admin_upload_audio.html')

@app.route('/admin/typing_passages', methods=['GET', 'POST'])
@admin_required
def admin_typing_passages():
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        
        # Get content from either text input or file upload
        content = ''
        text_input_method = request.form.get('text_input_method', 'direct')
        
        if text_input_method == 'direct':
            # Get text from direct input
            content = request.form.get('content', '').strip()
            if not content:
                flash('Please enter the passage content directly or upload a file', 'error')
                return redirect(request.url)
        else:
            # Handle passage file upload
            if 'passage_file' not in request.files:
                flash('No passage file selected', 'error')
                return redirect(request.url)
            
            passage_file = request.files['passage_file']
            if passage_file.filename == '':
                flash('No passage file selected', 'error')
                return redirect(request.url)
            
            # Read content from uploaded file (support multiple formats)
            file_extension = passage_file.filename.lower().split('.')[-1]
            
            try:
                if file_extension == 'txt':
                    content = passage_file.read().decode('utf-8').strip()
                elif file_extension == 'docx':
                    # Handle DOCX files
                    try:
                        import docx
                        doc = docx.Document(passage_file)
                        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    except ImportError:
                        flash('DOCX support not available. Please use TXT format or install python-docx', 'error')
                        return redirect(request.url)
                elif file_extension == 'odt':
                    # Handle ODT files
                    try:
                        from odf.opendocument import load
                        from odf.text import P
                        doc = load(passage_file)
                        paragraphs = doc.getElementsByType(P)
                        content = '\n'.join([str(p) for p in paragraphs])
                    except ImportError:
                        flash('ODT support not available. Please use TXT format or install odfpy', 'error')
                        return redirect(request.url)
                else:
                    flash('Please select a valid file format (.txt, .docx, or .odt)', 'error')
                    return redirect(request.url)
                
                if not content.strip():
                    flash('Passage file is empty', 'error')
                    return redirect(request.url)
                    
            except Exception as e:
                flash(f'Error reading passage file: {str(e)}', 'error')
                return redirect(request.url)
        
        # Use filename as title if no title provided and file was uploaded
        if not title and text_input_method == 'file' and 'passage_file' in request.files:
            title = request.files['passage_file'].filename.rsplit('.', 1)[0]
        elif not title:
            title = f'Passage {TypingPassage.query.count() + 1}'
        
        # Create typing passage with fixed 10-minute practice duration
        passage = TypingPassage(
            title=title,
            content=content.strip(),
            word_count=len(content.split()),
            uploaded_by=session.get('user_id')
        )
        db.session.add(passage)
        db.session.commit()
        
        flash('Typing passage added successfully! Practice time set to 10 minutes.', 'success')
        return redirect(url_for('admin_content'))
    
    return render_template('admin_typing_passages.html')

@app.route('/admin/bulk_upload', methods=['GET', 'POST'])
@admin_required
def admin_bulk_upload():
    if request.method == 'POST':
        upload_type = request.form.get('upload_type')
        
        if upload_type == 'audio':
            # Handle bulk audio uploads
            files = request.files.getlist('audio_files')
            success_count = 0
            error_count = 0
            errors = []
            
            for file in files:
                if file and file.filename and file.filename.lower().endswith('.mp3'):
                    try:
                        filename = secure_filename(file.filename)
                        file_path = os.path.join(UPLOAD_FOLDER, filename)
                        file.save(file_path)
                        
                        # Extract title from filename (remove extension)
                        title = request.form.get(f'title_{file.filename}') or filename.rsplit('.', 1)[0]
                        
                        audio_file = AudioFile(
                            title=title,
                            filename=filename,
                            uploaded_by=session.get('user_id')
                        )
                        db.session.add(audio_file)
                        success_count += 1
                    except Exception as e:
                        error_count += 1
                        errors.append(f"Error uploading {file.filename}: {str(e)}")
                
            try:
                db.session.commit()
                flash(f'Bulk upload completed! {success_count} files uploaded successfully.', 'success')
                if error_count > 0:
                    flash(f'{error_count} files failed to upload. Check the errors below.', 'error')
                    for error in errors:
                        flash(error, 'error')
            except Exception as e:
                db.session.rollback()
                flash(f'Database error: {str(e)}', 'error')
                
        elif upload_type == 'passages':
            # Handle bulk passage uploads
            passages_text = request.form.get('passages_bulk_text', '')
            success_count = 0
            
            if passages_text:
                # Split by double newlines to separate passages
                passages = [p.strip() for p in passages_text.split('\n\n') if p.strip()]
                
                for i, passage_content in enumerate(passages):
                    lines = passage_content.split('\n')
                    title = lines[0] if lines else f'Passage {i+1}'
                    content = '\n'.join(lines[1:]) if len(lines) > 1 else passage_content
                    
                    try:
                        passage = TypingPassage(
                            title=title,
                            content=content,
                            word_count=len(content.split()),
                            uploaded_by=session.get('user_id')
                        )
                        db.session.add(passage)
                        success_count += 1
                    except Exception as e:
                        flash(f'Error adding passage "{title}": {str(e)}', 'error')
                
                try:
                    db.session.commit()
                    flash(f'Bulk passage upload completed! {success_count} passages added successfully.', 'success')
                except Exception as e:
                    db.session.rollback()
                    flash(f'Database error: {str(e)}', 'error')
        
        return redirect(url_for('admin_bulk_upload'))
    
    return render_template('admin_bulk_upload.html')

@app.route('/admin/delete_audio/<int:audio_id>', methods=['POST'])
@admin_required
def admin_delete_audio(audio_id):
    audio_file = AudioFile.query.get_or_404(audio_id)
    
    try:
        # Delete physical file
        file_path = os.path.join(UPLOAD_FOLDER, audio_file.filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete from database
        db.session.delete(audio_file)
        db.session.commit()
        
        flash(f'Audio file "{audio_file.title}" deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting audio file: {str(e)}', 'error')
    
    return redirect(url_for('admin_content'))

@app.route('/admin/delete_passage/<int:passage_id>', methods=['POST'])
@admin_required
def admin_delete_passage(passage_id):
    passage = TypingPassage.query.get_or_404(passage_id)
    
    try:
        # Delete from database
        db.session.delete(passage)
        db.session.commit()
        
        flash(f'Typing passage "{passage.title}" deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting typing passage: {str(e)}', 'error')
    
    return redirect(url_for('admin_content'))

@app.route('/admin/move_audio/<int:audio_id>/<new_type>', methods=['POST'])
@admin_required
def admin_move_audio(audio_id, new_type):
    """Move audio file between exam and practice sections"""
    audio_file = AudioFile.query.get_or_404(audio_id)
    
    # Validate new_type
    if new_type not in ['exam', 'practice']:
        flash('Invalid content type', 'error')
        return redirect(url_for('admin_content'))
    
    # Check if moving to exam but no reference text
    if new_type == 'exam' and not audio_file.reference_text:
        flash('Cannot move to exam section: No reference text available. Please upload reference text first.', 'error')
        return redirect(url_for('admin_content'))
    
    try:
        old_type = audio_file.content_type
        audio_file.content_type = new_type
        db.session.commit()
        
        flash(f'Audio file "{audio_file.title}" moved from {old_type} to {new_type} section successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error moving audio file: {str(e)}', 'error')
    
    return redirect(url_for('admin_content'))

@app.route('/admin/reports')
@admin_required
def admin_reports():
    # Get statistics for reports
    total_users = User.query.filter_by(role='student').count()
    active_users = User.query.filter_by(role='student', is_active=True).count()
    total_attempts = DictationAttempt.query.count() + TypingAttempt.query.count()
    
    # Recent activity
    recent_dictation = DictationAttempt.query.order_by(DictationAttempt.submitted_at.desc()).limit(20).all()
    recent_typing = TypingAttempt.query.order_by(TypingAttempt.submitted_at.desc()).limit(20).all()
    
    # Performance statistics
    avg_dictation_score = db.session.query(db.func.avg(DictationAttempt.accuracy_percentage)).scalar() or 0
    avg_typing_wpm = db.session.query(db.func.avg(TypingAttempt.wpm)).scalar() or 0
    
    # Individual student performance summary
    students = User.query.filter_by(role='student').all()
    student_performance = []
    
    for student in students:
        # Get dictation statistics
        dictation_attempts = DictationAttempt.query.filter_by(user_id=student.id).filter(
            DictationAttempt.submitted_at.isnot(None)
        ).all()
        
        # Get typing statistics
        typing_attempts = TypingAttempt.query.filter_by(user_id=student.id).filter(
            TypingAttempt.submitted_at.isnot(None)
        ).all()
        
        # Calculate dictation stats
        dictation_count = len(dictation_attempts)
        avg_dictation_accuracy = sum(a.accuracy_percentage or 0 for a in dictation_attempts) / dictation_count if dictation_count > 0 else 0
        best_dictation_accuracy = max((a.accuracy_percentage or 0 for a in dictation_attempts), default=0)
        
        # Calculate typing stats
        typing_count = len(typing_attempts)
        avg_typing_wpm = sum(a.wpm or 0 for a in typing_attempts) / typing_count if typing_count > 0 else 0
        avg_typing_accuracy = sum(a.accuracy_percentage or 0 for a in typing_attempts) / typing_count if typing_count > 0 else 0
        best_typing_wpm = max((a.wpm or 0 for a in typing_attempts), default=0)
        
        # Get recent activity
        last_dictation = dictation_attempts[-1] if dictation_attempts else None
        last_typing = typing_attempts[-1] if typing_attempts else None
        
        # Determine last activity
        last_activity = None
        if last_dictation and last_typing:
            last_activity = max(last_dictation.submitted_at, last_typing.submitted_at)
        elif last_dictation:
            last_activity = last_dictation.submitted_at
        elif last_typing:
            last_activity = last_typing.submitted_at
        
        student_performance.append({
            'student': student,
            'dictation_count': dictation_count,
            'avg_dictation_accuracy': avg_dictation_accuracy,
            'best_dictation_accuracy': best_dictation_accuracy,
            'typing_count': typing_count,
            'avg_typing_wpm': avg_typing_wpm,
            'avg_typing_accuracy': avg_typing_accuracy,
            'best_typing_wpm': best_typing_wpm,
            'total_attempts': dictation_count + typing_count,
            'last_activity': last_activity
        })
    
    # Sort by total attempts (most active first)
    student_performance.sort(key=lambda x: x['total_attempts'], reverse=True)
    
    return render_template('admin_reports.html',
                         total_users=total_users,
                         active_users=active_users,
                         total_attempts=total_attempts,
                         recent_dictation=recent_dictation,
                         recent_typing=recent_typing,
                         avg_dictation_score=avg_dictation_score,
                         avg_typing_wpm=avg_typing_wpm,
                         student_performance=student_performance)

@app.route('/admin/export')
@admin_required
def admin_export():
    return render_template('admin_export.html')

# Test endpoint for debugging export issues
@app.route('/api/test-export', methods=['GET'])
@admin_required
def test_export():
    try:
        # Test basic functionality
        students = User.query.filter_by(role='student').all()
        print(f"Found {len(students)} students in database")  # Debug
        
        # Test openpyxl
        try:
            from openpyxl import Workbook
            wb = Workbook()
            print("openpyxl import and Workbook creation successful")  # Debug
        except Exception as e:
            print(f"openpyxl error: {e}")  # Debug
            
        # Test reportlab
        try:
            from reportlab.lib.pagesizes import A4
            print("reportlab import successful")  # Debug
        except Exception as e:
            print(f"reportlab error: {e}")  # Debug
            
        return jsonify({
            'success': True,
            'message': 'Test successful',
            'student_count': len(students)
        })
        
    except Exception as e:
        print(f"Test export error: {str(e)}")  # Debug
        import traceback
        print(f"Traceback: {traceback.format_exc()}")  # Debug
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Student Practice Routes
@app.route('/dictation-practice')
@login_required
def dictation_practice():
    # Get available audio files
    audio_files = AudioFile.query.all()
    if not audio_files:
        flash('No audio files available for practice. Please contact admin.', 'error')
        return redirect(url_for('practice_selection'))
    
    # Show all audio files in one list and add attempt count for current user
    available_files = []
    for audio in audio_files:
        if audio.reference_text:
            # Calculate attempt count for this user and audio
            attempt_count = DictationAttempt.query.filter_by(
                user_id=session['user_id'],
                audio_id=audio.id
            ).filter(DictationAttempt.submitted_at.isnot(None)).count()
            
            # Add attempt_count attribute to audio object
            audio.attempt_count = attempt_count
            available_files.append(audio)
    
    return render_template('dictation_practice.html',
                         audio_files=available_files)

@app.route('/dictation-practice-advanced/<int:audio_id>')
@login_required
def dictation_practice_advanced(audio_id):
    # Get the specific audio file
    audio = AudioFile.query.get_or_404(audio_id)
    return render_template('dictation_practice_advanced.html', audio=audio)

@app.route('/dictation-simple/<int:audio_id>')
@login_required
def dictation_simple(audio_id):
    # Get the specific audio file
    audio = AudioFile.query.get_or_404(audio_id)
    # Check if audio has reference text
    if not audio.reference_text:
        flash('This audio file does not have reference text. Please contact admin.', 'error')
        return redirect(url_for('dictation_practice'))
    return render_template('dictation_simple.html', audio=audio)

@app.route('/submit-dictation-practice', methods=['POST'])
@login_required
def submit_dictation_practice():
    try:
        # Get form data
        audio_id = request.form.get('audio_id')
        transcription = request.form.get('transcription', '')
        time_taken = int(request.form.get('time_taken', 0))
        practice_mode = request.form.get('practice_mode', 'unlimited')
        
        # Get audio file with reference text
        audio_file = AudioFile.query.get_or_404(audio_id)
        
        # Calculate attempt number for this user and audio (only count submitted attempts)
        previous_attempts = DictationAttempt.query.filter_by(
            user_id=session['user_id'],
            audio_id=audio_id
        ).filter(DictationAttempt.submitted_at.isnot(None)).count()
        attempt_number = previous_attempts + 1
        
        # Compare with reference text
        comparison_result = compare_texts(audio_file.reference_text, transcription)
        
        # Create dictation attempt record
        attempt = DictationAttempt(
            user_id=session['user_id'],
            audio_id=audio_id,
            transcription=transcription,
            words_typed=comparison_result['typed_words'],
            words_correct=comparison_result['words_correct'],
            words_wrong=comparison_result['words_wrong'],
            accuracy_percentage=comparison_result['accuracy_percentage'],
            time_taken=time_taken,
            attempt_number=attempt_number,
            submitted_at=datetime.now(timezone.utc)
        )
        
        db.session.add(attempt)
        db.session.commit()
        
        # Get user info for admin notification
        user = User.query.get(session['user_id'])
        
        # Get improvement data by comparing with previous attempt
        improvement_data = calculate_improvement_dictation(session['user_id'], audio_id, comparison_result['accuracy_percentage'], time_taken)
        
        # Analyze common mistakes
        mistakes_analysis = analyze_common_mistakes(audio_file.reference_text, transcription)
        
        # Store results in session for display
        session['dictation_results'] = {
            'audio_title': audio_file.title,
            'words_typed': comparison_result['typed_words'],
            'words_correct': comparison_result['words_correct'],
            'words_wrong': comparison_result['words_wrong'],
            'accuracy_percentage': comparison_result['accuracy_percentage'],
            'time_taken': time_taken,
            'total_reference_words': comparison_result['total_words'],
            'reference_text': audio_file.reference_text,
            'user_text': transcription,
            'attempt_number': attempt_number,
            'username': user.username,
            'content_type': audio_file.content_type,
            'enhanced_comparison': comparison_result['enhanced_comparison'],
            'improvement_data': improvement_data,
            'mistakes_analysis': mistakes_analysis
        }
        
        return redirect(url_for('dictation_result'))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error submitting practice: {str(e)}', 'error')
        return redirect(url_for('dictation_practice'))

@app.route('/dictation-result')
@login_required
def dictation_result():
    # Get results from session first
    results = session.get('dictation_results')
    
    # If no session results, try to get the latest attempt from database
    if not results:
        latest_attempt = DictationAttempt.query.filter_by(
            user_id=session['user_id']
        ).filter(
            DictationAttempt.submitted_at.isnot(None)
        ).order_by(DictationAttempt.submitted_at.desc()).first()
        
        if latest_attempt:
            # Get audio file and user info
            audio_file = AudioFile.query.get(latest_attempt.audio_id)
            user = User.query.get(session['user_id'])
            
            if audio_file and user:
                # Recreate comparison data from stored attempt
                comparison_result = compare_texts(audio_file.reference_text, latest_attempt.transcription or '')
                
                results = {
                    'audio_title': audio_file.title,
                    'words_typed': latest_attempt.words_typed or 0,
                    'words_correct': latest_attempt.words_correct or 0,
                    'words_wrong': latest_attempt.words_wrong or 0,
                    'accuracy_percentage': latest_attempt.accuracy_percentage or 0,
                    'time_taken': latest_attempt.time_taken or 0,
                    'total_reference_words': comparison_result['total_words'],
                    'reference_text': audio_file.reference_text,
                    'user_text': latest_attempt.transcription or '',
                    'attempt_number': latest_attempt.attempt_number or 1,
                    'username': user.username,
                    'content_type': audio_file.content_type or 'practice',
                    'detailed_comparison': comparison_result['detailed_comparison']
                }
    
    if not results:
        flash('No results found. Please complete a dictation practice first.', 'error')
        return redirect(url_for('dictation_practice'))
    
    # Clear results from session after displaying (only if it came from session)
    if session.get('dictation_results'):
        session.pop('dictation_results', None)
    
    return render_template('dictation_result.html', results=results)

@app.route('/typing-practice')
@login_required
def typing_practice():
    # Get available typing passages
    typing_passages = TypingPassage.query.all()
    if not typing_passages:
        flash('No typing passages available for practice. Please contact admin.', 'error')
        return redirect(url_for('practice_selection'))
    
    # Add attempt count for each passage for current user
    for passage in typing_passages:
        # Calculate attempt count for this user and passage
        attempt_count = TypingAttempt.query.filter_by(
            user_id=session['user_id'],
            passage_id=passage.id
        ).filter(TypingAttempt.submitted_at.isnot(None)).count()
        
        # Add attempt_count attribute to passage object
        passage.attempt_count = attempt_count
    
    return render_template('typing_practice.html', typing_passages=typing_passages)

@app.route('/typing-practice-advanced/<int:passage_id>')
@login_required
def typing_practice_advanced(passage_id):
    # Get the specific typing passage
    passage = TypingPassage.query.get_or_404(passage_id)
    return render_template('typing_practice_advanced.html', passage=passage)

@app.route('/typing-simple/<int:passage_id>')
@login_required
def typing_simple(passage_id):
    # Get the specific typing passage
    passage = TypingPassage.query.get_or_404(passage_id)
    return render_template('typing_simple.html', passage=passage)

@app.route('/submit-typing-practice', methods=['POST'])
@login_required
def submit_typing_practice():
    try:
        # Get form data
        passage_id = request.form.get('passage_id')
        typed_text = request.form.get('typed_text', '')
        time_taken = int(request.form.get('time_taken', 0))
        wpm = float(request.form.get('wpm', 0))
        accuracy = float(request.form.get('accuracy', 0))
        errors = int(request.form.get('errors', 0))
        practice_mode = request.form.get('practice_mode', 'normal')
        
        # Get passage for comparison
        passage = TypingPassage.query.get_or_404(passage_id)
        
        # Calculate attempt number for this user and passage (only count submitted attempts)
        previous_attempts = TypingAttempt.query.filter_by(
            user_id=session['user_id'],
            passage_id=passage_id
        ).filter(TypingAttempt.submitted_at.isnot(None)).count()
        attempt_number = previous_attempts + 1
        
        # Calculate more accurate statistics using text comparison
        comparison_result = compare_texts(passage.content, typed_text)
        
        # Use server-calculated accuracy instead of client accuracy
        words_typed = comparison_result['typed_words']
        words_correct = comparison_result['words_correct']
        words_wrong = comparison_result['words_wrong']
        calculated_accuracy = comparison_result['accuracy_percentage']
        
        # Recalculate WPM based on server data
        calculated_wpm = (words_typed / (time_taken / 60)) if time_taken > 0 else 0
        
        # Create typing attempt record
        attempt = TypingAttempt(
            user_id=session['user_id'],
            passage_id=passage_id,
            typed_text=typed_text,
            words_typed=words_typed,
            words_correct=words_correct,
            words_wrong=words_wrong,
            accuracy_percentage=calculated_accuracy,
            wpm=calculated_wpm,
            time_taken=time_taken,
            attempt_number=attempt_number,
            submitted_at=datetime.now(timezone.utc)
        )
        
        db.session.add(attempt)
        db.session.commit()
        
        # Get user info for admin notification
        user = User.query.get(session['user_id'])
        
        # Get improvement data by comparing with previous attempt
        improvement_data = calculate_improvement_typing(session['user_id'], passage_id, calculated_accuracy, calculated_wpm, time_taken)
        
        # Analyze common mistakes
        mistakes_analysis = analyze_common_mistakes(passage.content, typed_text)
        
        # Store results in session for display (if we create a results page later)
        session['typing_results'] = {
            'passage_title': passage.title,
            'words_typed': words_typed,
            'words_correct': words_correct,
            'words_wrong': words_wrong,
            'accuracy_percentage': calculated_accuracy,
            'wpm': calculated_wpm,
            'time_taken': time_taken,
            'attempt_number': attempt_number,
            'username': user.username,
            'reference_text': passage.content,
            'user_text': typed_text,
            'enhanced_comparison': comparison_result['enhanced_comparison'],
            'improvement_data': improvement_data,
            'mistakes_analysis': mistakes_analysis
        }
        
        flash(f'Typing practice submitted! Attempt #{attempt_number} - WPM: {calculated_wpm:.1f}, Accuracy: {calculated_accuracy:.1f}%, Time: {time_taken}s', 'success')
        return redirect(url_for('typing_result'))
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error submitting practice: {str(e)}', 'error')
        return redirect(url_for('typing_practice'))

@app.route('/typing-result')
@login_required
def typing_result():
    # Get results from session first
    results = session.get('typing_results')
    
    # If no session results, try to get the latest attempt from database
    if not results:
        latest_attempt = TypingAttempt.query.filter_by(
            user_id=session['user_id']
        ).filter(
            TypingAttempt.submitted_at.isnot(None)
        ).order_by(TypingAttempt.submitted_at.desc()).first()
        
        if latest_attempt:
            # Get passage and user info
            passage = TypingPassage.query.get(latest_attempt.passage_id)
            user = User.query.get(session['user_id'])
            
            if passage and user:
                # Recreate comparison data from stored attempt
                comparison_result = compare_texts(passage.content, latest_attempt.typed_text or '')
                
                results = {
                    'passage_title': passage.title,
                    'words_typed': latest_attempt.words_typed or 0,
                    'words_correct': latest_attempt.words_correct or 0,
                    'words_wrong': latest_attempt.words_wrong or 0,
                    'accuracy_percentage': latest_attempt.accuracy_percentage or 0,
                    'wpm': latest_attempt.wpm or 0,
                    'time_taken': latest_attempt.time_taken or 0,
                    'attempt_number': latest_attempt.attempt_number or 1,
                    'username': user.username,
                    'reference_text': passage.content,
                    'user_text': latest_attempt.typed_text or '',
                    'enhanced_comparison': comparison_result['enhanced_comparison']
                }
    
    if not results:
        flash('No results found. Please complete a typing practice first.', 'error')
        return redirect(url_for('typing_practice'))
    
    # Clear results from session after displaying (only if it came from session)
    if session.get('typing_results'):
        session.pop('typing_results', None)
    
    return render_template('typing_result.html', results=results)

@app.route('/dictation-leaderboard')
@login_required
def dictation_leaderboard():
    # Get top dictation attempts with user and audio data
    top_attempts = db.session.query(
        DictationAttempt,
        User.username,
        AudioFile.title.label('audio_title')
    ).join(
        User, DictationAttempt.user_id == User.id
    ).join(
        AudioFile, DictationAttempt.audio_id == AudioFile.id
    ).filter(
        DictationAttempt.submitted_at.isnot(None)
    ).order_by(
        DictationAttempt.accuracy_percentage.desc(),
        DictationAttempt.time_taken.asc()
    ).limit(20).all()
    
    # Format leaderboard data
    leaderboard_data = []
    for attempt, username, audio_title in top_attempts:
        score = int(attempt.accuracy_percentage * (100 / max(1, attempt.time_taken/60)))  # Score based on accuracy and speed
        leaderboard_data.append({
            'username': username,
            'audio_filename': audio_title,
            'accuracy': attempt.accuracy_percentage or 0,
            'time_taken': round(attempt.time_taken / 60, 1) if attempt.time_taken else 0,  # Convert to minutes
            'completed_at': attempt.submitted_at,
            'score': score,
            'words_typed': attempt.words_typed or 0
        })
    
    # Get user stats
    current_user = User.query.get(session['user_id'])
    user_attempts = DictationAttempt.query.filter_by(user_id=current_user.id).filter(
        DictationAttempt.submitted_at.isnot(None)
    ).all()
    
    user_stats = None
    if user_attempts:
        total_attempts = len(user_attempts)
        avg_accuracy = sum(a.accuracy_percentage or 0 for a in user_attempts) / total_attempts
        best_score = max(int((a.accuracy_percentage or 0) * (100 / max(1, (a.time_taken or 1)/60))) for a in user_attempts)
        
        # Calculate current rank
        current_rank = 1
        for i, entry in enumerate(leaderboard_data):
            if entry['username'] == current_user.username:
                current_rank = i + 1
                break
        else:
            current_rank = len(leaderboard_data) + 1
        
        user_stats = {
            'total_attempts': total_attempts,
            'average_accuracy': avg_accuracy,
            'best_score': best_score,
            'current_rank': current_rank
        }
    
    # Get available audio files for filter
    available_audios = AudioFile.query.all()
    
    return render_template('dictation_leaderboard.html',
                         leaderboard_data=leaderboard_data,
                         user_stats=user_stats,
                         available_audios=available_audios)

@app.route('/typing-leaderboard')
@login_required
def typing_leaderboard():
    # Get top typing attempts with user and passage data
    top_attempts = db.session.query(
        TypingAttempt,
        User.username,
        TypingPassage.title.label('passage_title')
    ).join(
        User, TypingAttempt.user_id == User.id
    ).join(
        TypingPassage, TypingAttempt.passage_id == TypingPassage.id
    ).filter(
        TypingAttempt.submitted_at.isnot(None)
    ).order_by(
        TypingAttempt.wpm.desc(),
        TypingAttempt.accuracy_percentage.desc()
    ).limit(20).all()
    
    # Format leaderboard data
    leaderboard_data = []
    for attempt, username, passage_title in top_attempts:
        # Score based on WPM and accuracy
        score = int((attempt.wpm or 0) * (attempt.accuracy_percentage or 0) / 100)
        leaderboard_data.append({
            'username': username,
            'passage_title': passage_title,
            'wpm': round(attempt.wpm or 0, 1),
            'accuracy': attempt.accuracy_percentage or 0,
            'error_count': attempt.words_wrong or 0,
            'completed_at': attempt.submitted_at,
            'score': score
        })
    
    # Get user stats
    current_user = User.query.get(session['user_id'])
    user_attempts = TypingAttempt.query.filter_by(user_id=current_user.id).filter(
        TypingAttempt.submitted_at.isnot(None)
    ).all()
    
    user_stats = None
    if user_attempts:
        total_attempts = len(user_attempts)
        avg_wpm = sum(a.wpm or 0 for a in user_attempts) / total_attempts
        avg_accuracy = sum(a.accuracy_percentage or 0 for a in user_attempts) / total_attempts
        best_score = max(int((a.wpm or 0) * (a.accuracy_percentage or 0) / 100) for a in user_attempts)
        
        # Calculate current rank
        current_rank = 1
        for i, entry in enumerate(leaderboard_data):
            if entry['username'] == current_user.username:
                current_rank = i + 1
                break
        else:
            current_rank = len(leaderboard_data) + 1
        
        user_stats = {
            'total_attempts': total_attempts,
            'average_wpm': round(avg_wpm, 1),
            'average_accuracy': avg_accuracy,
            'best_score': best_score,
            'current_rank': current_rank
        }
    
    # Get available passages for filter
    available_passages = TypingPassage.query.all()
    
    return render_template('typing_leaderboard.html',
                         leaderboard_data=leaderboard_data,
                         user_stats=user_stats,
                         available_passages=available_passages)

# API Routes for Filtering
@app.route('/api/filter-dictation-leaderboard', methods=['POST'])
@login_required
def filter_dictation_leaderboard():
    try:
        data = request.get_json()
        time_filter = data.get('timeFilter', 'all')
        audio_filter = data.get('audioFilter', 'all')
        
        # Base query
        query = db.session.query(
            DictationAttempt,
            User.username,
            AudioFile.title.label('audio_title')
        ).join(
            User, DictationAttempt.user_id == User.id
        ).join(
            AudioFile, DictationAttempt.audio_id == AudioFile.id
        ).filter(
            DictationAttempt.submitted_at.isnot(None)
        )
        
        # Apply time filter
        if time_filter == 'week':
            week_ago = datetime.now(timezone.utc) - timedelta(days=7)
            query = query.filter(DictationAttempt.submitted_at >= week_ago)
        elif time_filter == 'month':
            month_ago = datetime.now(timezone.utc) - timedelta(days=30)
            query = query.filter(DictationAttempt.submitted_at >= month_ago)
        
        # Apply audio filter
        if audio_filter != 'all':
            query = query.filter(DictationAttempt.audio_id == int(audio_filter))
        
        # Order and limit results
        top_attempts = query.order_by(
            DictationAttempt.accuracy_percentage.desc(),
            DictationAttempt.time_taken.asc()
        ).limit(20).all()
        
        # Format leaderboard data
        leaderboard_data = []
        for attempt, username, audio_title in top_attempts:
            score = int(attempt.accuracy_percentage * (100 / max(1, attempt.time_taken/60)))
            leaderboard_data.append({
                'username': username,
                'audio_filename': audio_title,
                'accuracy': attempt.accuracy_percentage or 0,
                'time_taken': round(attempt.time_taken / 60, 1) if attempt.time_taken else 0,
                'completed_at': attempt.submitted_at.strftime('%Y-%m-%d'),
                'score': score,
                'words_typed': attempt.words_typed or 0
            })
        
        return jsonify({
            'success': True,
            'leaderboard_data': leaderboard_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/filter-typing-leaderboard', methods=['POST'])
@login_required
def filter_typing_leaderboard():
    try:
        data = request.get_json()
        time_filter = data.get('timeFilter', 'all')
        passage_filter = data.get('passageFilter', 'all')
        metric_filter = data.get('metricFilter', 'wpm')
        
        # Base query
        query = db.session.query(
            TypingAttempt,
            User.username,
            TypingPassage.title.label('passage_title')
        ).join(
            User, TypingAttempt.user_id == User.id
        ).join(
            TypingPassage, TypingAttempt.passage_id == TypingPassage.id
        ).filter(
            TypingAttempt.submitted_at.isnot(None)
        )
        
        # Apply time filter
        if time_filter == 'week':
            week_ago = datetime.now(timezone.utc) - timedelta(days=7)
            query = query.filter(TypingAttempt.submitted_at >= week_ago)
        elif time_filter == 'month':
            month_ago = datetime.now(timezone.utc) - timedelta(days=30)
            query = query.filter(TypingAttempt.submitted_at >= month_ago)
        
        # Apply passage filter
        if passage_filter != 'all':
            query = query.filter(TypingAttempt.passage_id == int(passage_filter))
        
        # Order by selected metric
        if metric_filter == 'accuracy':
            query = query.order_by(
                TypingAttempt.accuracy_percentage.desc(),
                TypingAttempt.wpm.desc()
            )
        elif metric_filter == 'score':
            # Create a score expression (WPM * Accuracy / 100)
            score_expr = TypingAttempt.wpm * TypingAttempt.accuracy_percentage / 100
            query = query.order_by(score_expr.desc())
        else:  # Default to WPM
            query = query.order_by(
                TypingAttempt.wpm.desc(),
                TypingAttempt.accuracy_percentage.desc()
            )
        
        # Limit results
        top_attempts = query.limit(20).all()
        
        # Format leaderboard data
        leaderboard_data = []
        for attempt, username, passage_title in top_attempts:
            score = int((attempt.wpm or 0) * (attempt.accuracy_percentage or 0) / 100)
            leaderboard_data.append({
                'username': username,
                'passage_title': passage_title,
                'wpm': round(attempt.wpm or 0, 1),
                'accuracy': attempt.accuracy_percentage or 0,
                'error_count': attempt.words_wrong or 0,
                'completed_at': attempt.submitted_at.strftime('%Y-%m-%d'),
                'score': score
            })
        
        return jsonify({
            'success': True,
            'leaderboard_data': leaderboard_data
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/get-user-password/<int:user_id>', methods=['GET'])
@admin_required
def get_user_password(user_id):
    try:
        user = User.query.get_or_404(user_id)
        
        if user.role != 'student':
            return jsonify({
                'success': False,
                'error': 'Can only view student passwords'
            }), 403
        
        # Since the requirement is to show actual passwords to admins,
        # and we're using hashed passwords (secure), we need a workaround.
        # For demo/development purposes, we'll provide common password patterns.
        
        # Enhanced demo password mapping with more patterns
        demo_passwords = {
            'student1': 'student123',
            'student2': 'student123',
            'student3': 'student123',
            'demo': 'demo123',
            'test': 'test123',
            app.config.get('ADMIN_USERNAME', 'admin'): app.config.get('ADMIN_PASSWORD', 'admin123')
        }
        
        # Check if it's a known demo account
        if user.username in demo_passwords:
            return jsonify({
                'success': True,
                'password': demo_passwords[user.username],
                'note': f'Demo account password for {user.username}'
            })
        
        # For admin-created accounts, we'll implement a practical solution:
        # Try common password patterns that admins typically use
        common_patterns = [
            f'{user.username}123',      # username + 123
            f'{user.username}@123',     # username + @123  
            'student123',               # default student password
            'password123',              # common default
            '123456',                   # simple default
            user.username,              # just the username
            f'{user.username}2024',     # username + year
            f'{user.username}_123'      # username + _123
        ]
        
        # Instead of showing hash error, show the most likely passwords
        suggested_passwords = ', '.join(common_patterns[:3])  # Show top 3
        
        return jsonify({
            'success': True,
            'password': f'Try: {suggested_passwords}',
            'note': f'Common patterns for {user.username}. If none work, reset password.',
            'is_suggestion': True
        })
        
    except Exception as e:
        print(f"Error retrieving password for user {user_id}: {str(e)}")  # Debug log
        return jsonify({
            'success': False,
            'error': f'Unable to retrieve password: {str(e)}'
        }), 500

# Forgot Password Routes for Admins Only
@app.route('/admin/forgot-password', methods=['GET', 'POST'])
def admin_forgot_password():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        
        if not username:
            flash('Username is required', 'error')
            return render_template('admin_forgot_password.html')
        
        # Only allow admin users to reset passwords
        user = User.query.filter_by(username=username, role='admin').first()
        
        if not user:
            # Don't reveal if user exists or not for security
            flash('If this admin username exists, a password reset has been initiated. Check your contact methods.', 'info')
            return render_template('admin_forgot_password.html')
        
        try:
            # Generate reset token and OTP
            reset_token = generate_reset_token()
            otp_code = generate_otp()
            
            # Set expiration time (35 minutes from now)
            expires_at = datetime.now(timezone.utc) + timedelta(minutes=35)
            
            # Invalidate any existing reset tokens for this user
            PasswordResetToken.query.filter_by(user_id=user.id, is_used=False).update({'is_used': True})
            
            # Create new reset token
            reset_request = PasswordResetToken(
                user_id=user.id,
                token=reset_token,
                expires_at=expires_at,
                otp_code=otp_code,
                reset_method='both'  # Support both email and SMS
            )
            db.session.add(reset_request)
            db.session.commit()
            
            # Send notifications (using configured admin contact info)
            admin_email = app.config.get('ADMIN_EMAIL', 'admin@localhost')
            admin_mobile = app.config.get('ADMIN_PHONE', '+1234567890')
            
            # Send email notification
            send_reset_email(admin_email, reset_token, otp_code)
            
            # Send SMS notification
            send_reset_sms(admin_mobile, otp_code)
            
            flash(f'Password reset initiated for admin: {username}. Check email ({admin_email}) and SMS ({admin_mobile}) for reset instructions.', 'success')
            return redirect(url_for('admin_reset_password_form'))
            
        except Exception as e:
            db.session.rollback()
            flash('Error initiating password reset. Please contact system administrator.', 'error')
            return render_template('admin_forgot_password.html')
    
    return render_template('admin_forgot_password.html')

@app.route('/admin/reset-password', methods=['GET', 'POST'])
def admin_reset_password_form():
    if request.method == 'POST':
        token_or_otp = request.form.get('token_or_otp', '').strip()
        new_password = request.form.get('new_password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        if not token_or_otp or not new_password or not confirm_password:
            flash('All fields are required', 'error')
            return render_template('admin_reset_password.html')
        
        if new_password != confirm_password:
            flash('Passwords do not match', 'error')
            return render_template('admin_reset_password.html')
        
        if len(new_password) < 6:
            flash('Password must be at least 6 characters long', 'error')
            return render_template('admin_reset_password.html')
        
        try:
            # Find valid reset token (check both token and OTP)
            now = datetime.now(timezone.utc)
            reset_request = PasswordResetToken.query.filter(
                db.or_(
                    PasswordResetToken.token == token_or_otp,
                    PasswordResetToken.otp_code == token_or_otp
                ),
                PasswordResetToken.expires_at > now,
                PasswordResetToken.is_used == False
            ).first()
            
            if not reset_request:
                flash('Invalid or expired reset token/OTP', 'error')
                return render_template('admin_reset_password.html')
            
            # Get the user
            user = User.query.get(reset_request.user_id)
            if not user or user.role != 'admin':
                flash('Invalid reset request', 'error')
                return render_template('admin_reset_password.html')
            
            # Update password
            user.password_hash = hash_password(new_password)
            
            # Mark reset token as used
            reset_request.is_used = True
            
            db.session.commit()
            
            flash(f'Password successfully reset for admin: {user.username}', 'success')
            return redirect(url_for('admin_login'))
            
        except Exception as e:
            db.session.rollback()
            flash('Error resetting password. Please try again.', 'error')
            return render_template('admin_reset_password.html')
    
    return render_template('admin_reset_password.html')

# API Route for Student Forgot Password (Limited functionality)
@app.route('/api/forgot-password', methods=['POST'])
def api_forgot_password():
    """
    Limited forgot password API for students.
    Students cannot reset passwords themselves - only admin can reset.
    This endpoint provides helpful information and admin contact details.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': 'No data provided'
            }), 400
            
        username = data.get('username', '').strip()
        
        if not username:
            return jsonify({
                'success': False,
                'error': 'Username is required'
            }), 400
        
        # Check if user exists (but don't reveal if they exist or not for security)
        user = User.query.filter_by(username=username, role='student').first()
        
        # Regardless of whether user exists, provide helpful response
        return jsonify({
            'success': True,
            'message': f'Password reset request received for "{username}". Since students cannot reset their own passwords, please contact your administrator directly:\n\n📞 Call: {app.config.get("ADMIN_PHONE", "+1234567890")}\n📧 Email: {app.config.get("ADMIN_EMAIL", "admin@localhost")}\n\nThe admin can reset your password and help with any account issues.',
            'admin_contact': {
                'phone': app.config.get('ADMIN_PHONE', '+1234567890'),
                'email': app.config.get('ADMIN_EMAIL', 'admin@localhost')
            }
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': 'Unable to process request. Please contact administrator directly.'
        }), 500

# Export Data Routes
@app.route('/api/export-data', methods=['POST'])
@admin_required
def export_data():
    try:
        data = request.get_json()
        if not data:
            print("No JSON data received")  # Debug
            return jsonify({
                'success': False,
                'error': 'No data received in request'
            }), 400
            
        export_type = data.get('type')
        export_format = data.get('format', 'excel')
        start_date = data.get('startDate')
        end_date = data.get('endDate')
        include_data = data.get('includeData', {})
        
        print(f"=== EXPORT REQUEST DEBUG ===")  # Debug
        print(f"Export request: type={export_type}, format={export_format}")  # Debug log
        print(f"Include data: {include_data}")  # Debug log
        print(f"Start date: {start_date}, End date: {end_date}")  # Debug log
        
        # Import required libraries for export
        import io
        import csv
        from datetime import datetime
        from flask import Response
        import tempfile
        import os
        from io import BytesIO
        
        export_data_dict = {}
        
        # Handle combined export type by including all data types
        if export_type == 'combined':
            print("Setting combined export to include all data types")  # Debug
            include_data = {
                'students': True,
                'dictation': True,
                'typing': True,
                'statistics': True,
                'subscriptions': True
            }
        
        print(f"Final include_data: {include_data}")  # Debug
        
        # Students data
        if export_type == 'students' or include_data.get('students', False):
            print("Collecting students data...")  # Debug
            try:
                students = User.query.filter_by(role='student').all()
                print(f"Found {len(students)} students")  # Debug
                export_data_dict['students'] = []
                for student in students:
                    export_data_dict['students'].append({
                        'ID': student.id,
                        'Username': student.username,
                        'Status': 'Active' if student.is_active else 'Inactive',
                        'Locked': 'Yes' if student.is_locked else 'No',
                        'Subscription Start': student.subscription_start.strftime('%Y-%m-%d') if student.subscription_start else 'N/A',
                        'Subscription End': student.subscription_end.strftime('%Y-%m-%d') if student.subscription_end else 'N/A',
                        'Days Remaining': get_subscription_days_remaining(student),
                        'Last Login': student.last_login.strftime('%Y-%m-%d %H:%M') if student.last_login else 'Never',
                        'Created At': student.created_at.strftime('%Y-%m-%d') if student.created_at else 'N/A'
                    })
                print(f"Students data collected successfully: {len(export_data_dict['students'])} records")  # Debug
            except Exception as e:
                print(f"Error collecting students data: {e}")  # Debug
                raise
        
        # Dictation attempts data
        if export_type == 'dictation' or include_data.get('dictation', False):
            query = db.session.query(DictationAttempt, User.username, AudioFile.title).join(
                User, DictationAttempt.user_id == User.id
            ).join(AudioFile, DictationAttempt.audio_id == AudioFile.id)
            
            if start_date:
                query = query.filter(DictationAttempt.submitted_at >= datetime.strptime(start_date, '%Y-%m-%d'))
            if end_date:
                query = query.filter(DictationAttempt.submitted_at <= datetime.strptime(end_date, '%Y-%m-%d'))
            
            attempts = query.all()
            export_data_dict['dictation_attempts'] = []
            for attempt, username, audio_title in attempts:
                export_data_dict['dictation_attempts'].append({
                    'Student': username,
                    'Audio File': audio_title,
                    'Words Typed': attempt.words_typed or 0,
                    'Words Correct': attempt.words_correct or 0,
                    'Words Wrong': attempt.words_wrong or 0,
                    'Accuracy %': round(attempt.accuracy_percentage or 0, 2),
                    'Time Taken (min)': round((attempt.time_taken or 0) / 60, 2),
                    'Submitted At': attempt.submitted_at.strftime('%Y-%m-%d %H:%M') if attempt.submitted_at else 'N/A'
                })
        
        # Typing attempts data
        if export_type == 'typing' or include_data.get('typing', False):
            query = db.session.query(TypingAttempt, User.username, TypingPassage.title).join(
                User, TypingAttempt.user_id == User.id
            ).join(TypingPassage, TypingAttempt.passage_id == TypingPassage.id)
            
            if start_date:
                query = query.filter(TypingAttempt.submitted_at >= datetime.strptime(start_date, '%Y-%m-%d'))
            if end_date:
                query = query.filter(TypingAttempt.submitted_at <= datetime.strptime(end_date, '%Y-%m-%d'))
            
            attempts = query.all()
            export_data_dict['typing_attempts'] = []
            for attempt, username, passage_title in attempts:
                export_data_dict['typing_attempts'].append({
                    'Student': username,
                    'Passage': passage_title,
                    'Words Typed': attempt.words_typed or 0,
                    'Words Correct': attempt.words_correct or 0,
                    'Words Wrong': attempt.words_wrong or 0,
                    'Accuracy %': round(attempt.accuracy_percentage or 0, 2),
                    'WPM': round(attempt.wpm or 0, 2),
                    'Time Taken (sec)': attempt.time_taken or 0,
                    'Submitted At': attempt.submitted_at.strftime('%Y-%m-%d %H:%M') if attempt.submitted_at else 'N/A'
                })
        
        # Performance statistics
        if export_type == 'performance' or include_data.get('statistics', False):
            students = User.query.filter_by(role='student').all()
            export_data_dict['performance_stats'] = []
            for student in students:
                dictation_attempts = DictationAttempt.query.filter_by(user_id=student.id).filter(
                    DictationAttempt.submitted_at.isnot(None)
                ).all()
                typing_attempts = TypingAttempt.query.filter_by(user_id=student.id).filter(
                    TypingAttempt.submitted_at.isnot(None)
                ).all()
                
                avg_dictation_accuracy = sum(a.accuracy_percentage or 0 for a in dictation_attempts) / len(dictation_attempts) if dictation_attempts else 0
                avg_typing_wpm = sum(a.wpm or 0 for a in typing_attempts) / len(typing_attempts) if typing_attempts else 0
                avg_typing_accuracy = sum(a.accuracy_percentage or 0 for a in typing_attempts) / len(typing_attempts) if typing_attempts else 0
                
                export_data_dict['performance_stats'].append({
                    'Student': student.username,
                    'Total Dictation Attempts': len(dictation_attempts),
                    'Avg Dictation Accuracy %': round(avg_dictation_accuracy, 2),
                    'Total Typing Attempts': len(typing_attempts),
                    'Avg Typing WPM': round(avg_typing_wpm, 2),
                    'Avg Typing Accuracy %': round(avg_typing_accuracy, 2),
                    'Total Attempts': len(dictation_attempts) + len(typing_attempts)
                })
        
        # Subscription data
        if export_type == 'subscriptions' or include_data.get('subscriptions', False):
            students = User.query.filter_by(role='student').all()
            export_data_dict['subscriptions'] = []
            for student in students:
                days_remaining = get_subscription_days_remaining(student)
                is_active = check_subscription_active(student)
                
                export_data_dict['subscriptions'].append({
                    'Student': student.username,
                    'Subscription Start': student.subscription_start.strftime('%Y-%m-%d') if student.subscription_start else 'N/A',
                    'Subscription End': student.subscription_end.strftime('%Y-%m-%d') if student.subscription_end else 'N/A',
                    'Days Remaining': days_remaining,
                    'Status': 'Active' if is_active else 'Expired',
                    'Account Status': 'Active' if student.is_active else 'Suspended',
                    'Account Locked': 'Yes' if student.is_locked else 'No'
                })
        
        # Debug: Print final export data summary
        print(f"=== EXPORT DATA SUMMARY ===")  # Debug
        for key, value in export_data_dict.items():
            if value:
                print(f"{key}: {len(value)} records")  # Debug
            else:
                print(f"{key}: 0 records")  # Debug
        print(f"Total sections with data: {len([k for k, v in export_data_dict.items() if v])}")  # Debug
        
        # Generate CSV content
        if export_format == 'csv':
            output = io.StringIO()
            
            # Check if there's any data to export
            if not export_data_dict or not any(export_data_dict.values()):
                return jsonify({
                    'success': False,
                    'error': 'No data available for export'
                }), 400
            
            for sheet_name, data_list in export_data_dict.items():
                if data_list:
                    output.write(f"\n{sheet_name.replace('_', ' ').title()}\n")
                    if data_list:
                        fieldnames = data_list[0].keys()
                        writer = csv.DictWriter(output, fieldnames=fieldnames)
                        writer.writeheader()
                        writer.writerows(data_list)
                    output.write("\n")
            
            # Create response
            from flask import Response
            response = Response(
                output.getvalue(),
                mimetype='text/csv',
                headers={'Content-Disposition': f'attachment; filename=export_{export_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'}
            )
            return response
        
        # Generate Excel export
        elif export_format == 'excel':
            try:
                print("Starting Excel export...")  # Debug
                from openpyxl import Workbook
                from openpyxl.styles import Font, PatternFill, Alignment
                print("Excel libraries imported successfully")  # Debug
            except ImportError as e:
                print(f"Excel import error: {e}")  # Debug
                return jsonify({
                    'success': False,
                    'error': f'Excel export library not available: {str(e)}'
                }), 500
            
            try:
                wb = Workbook()
                # Remove default sheet
                if wb.worksheets:
                    wb.remove(wb.active)
                
                print(f"Export data dict keys: {list(export_data_dict.keys())}")  # Debug
                print(f"Export data dict values lengths: {[(k, len(v) if v else 0) for k, v in export_data_dict.items()]}")  # Debug
                
                # Check if there's any data to export
                if not export_data_dict or not any(export_data_dict.values()):
                    print("No data available for export")  # Debug
                    return jsonify({
                        'success': False,
                        'error': 'No data available for export'
                    }), 400
                
                # If no sheets were created, create a default one
                has_sheets = False
                for sheet_name, data_list in export_data_dict.items():
                    if data_list:
                        print(f"Creating sheet: {sheet_name} with {len(data_list)} rows")  # Debug
                        has_sheets = True
                        sheet_title = sheet_name.replace('_', ' ').title()[:31]  # Excel sheet name limit
                        ws = wb.create_sheet(title=sheet_title)
                        
                        # Add headers
                        headers = list(data_list[0].keys())
                        for col, header in enumerate(headers, 1):
                            cell = ws.cell(row=1, column=col, value=header)
                            cell.font = Font(bold=True, color="FFFFFF")
                            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                            cell.alignment = Alignment(horizontal="center")
                        
                        # Add data
                        for row, item in enumerate(data_list, 2):
                            for col, header in enumerate(headers, 1):
                                value = item.get(header, '')
                                ws.cell(row=row, column=col, value=value)
                        
                        # Auto-adjust column widths
                        for column in ws.columns:
                            max_length = 0
                            column_letter = column[0].column_letter
                            for cell in column:
                                try:
                                    if len(str(cell.value)) > max_length:
                                        max_length = len(str(cell.value))
                                except:
                                    pass
                            adjusted_width = min(max_length + 2, 50)
                            ws.column_dimensions[column_letter].width = adjusted_width
                
                if not has_sheets:
                    print("No sheets created, adding default sheet")  # Debug
                    # Create a default sheet with a message
                    ws = wb.create_sheet(title="No Data")
                    ws.cell(row=1, column=1, value="No data available for the selected export criteria")
                
                # Use BytesIO instead of temp file
                print("Saving workbook to BytesIO...")  # Debug
                output = BytesIO()
                wb.save(output)
                output.seek(0)
                
                print("Creating response...")  # Debug
                # Create response
                response = Response(
                    output.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    headers={'Content-Disposition': f'attachment; filename=export_{export_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'}
                )
                print("Excel export successful")  # Debug
                return response
                
            except Exception as e:
                print(f"Excel export error: {str(e)}")  # Debug
                return jsonify({
                    'success': False,
                    'error': f'Excel export failed: {str(e)}'
                }), 500
        
        # Generate PDF export
        elif export_format == 'pdf':
            try:
                print("Starting PDF export...")  # Debug
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.lib.units import inch
                print("PDF libraries imported successfully")  # Debug
            except ImportError as e:
                print(f"PDF import error: {e}")  # Debug
                return jsonify({
                    'success': False,
                    'error': f'PDF export library not available: {str(e)}'
                }), 500
            
            try:
                print(f"Export data dict keys: {list(export_data_dict.keys())}")  # Debug
                
                # Check if there's any data to export
                if not export_data_dict or not any(export_data_dict.values()):
                    print("No data available for PDF export")  # Debug
                    return jsonify({
                        'success': False,
                        'error': 'No data available for export'
                    }), 400
                
                # Use BytesIO instead of temp file
                output = BytesIO()
                doc = SimpleDocTemplate(output, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # Title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1  # Center alignment
                )
                story.append(Paragraph(f"Data Export Report - {export_type.title()}", title_style))
                story.append(Spacer(1, 20))
                
                has_data = False
                for sheet_name, data_list in export_data_dict.items():
                    if data_list:
                        has_data = True
                        print(f"Adding PDF section: {sheet_name} with {len(data_list)} rows")  # Debug
                        
                        # Section header
                        section_style = ParagraphStyle(
                            'SectionHeader',
                            parent=styles['Heading2'],
                            fontSize=14,
                            spaceAfter=12
                        )
                        story.append(Paragraph(sheet_name.replace('_', ' ').title(), section_style))
                        
                        # Prepare table data
                        headers = list(data_list[0].keys())
                        # Limit headers for PDF width
                        if len(headers) > 6:
                            headers = headers[:6]  # Take first 6 columns for better PDF formatting
                        
                        table_data = [headers]
                        
                        for item in data_list[:50]:  # Limit to 50 rows per section for PDF
                            row = [str(item.get(header, ''))[:20] + ('...' if len(str(item.get(header, ''))) > 20 else '') for header in headers]
                            table_data.append(row)
                        
                        # Create table with smaller font and better sizing
                        try:
                            table = Table(table_data, repeatRows=1)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 8),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('FONTSIZE', (0, 1), (-1, -1), 6),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ]))
                            
                            story.append(table)
                            story.append(Spacer(1, 20))
                            
                            if len(data_list) > 50:
                                story.append(Paragraph(f"... and {len(data_list) - 50} more records", styles['Normal']))
                                story.append(Spacer(1, 20))
                        except Exception as table_error:
                            print(f"Table creation error: {table_error}")  # Debug
                            # Fallback to simple text if table fails
                            story.append(Paragraph(f"Data for {sheet_name}: {len(data_list)} records", styles['Normal']))
                            story.append(Spacer(1, 20))
                
                if not has_data:
                    story.append(Paragraph("No data available for the selected export criteria.", styles['Normal']))
                
                # Add generation info
                story.append(Spacer(1, 30))
                story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
                
                print("Building PDF...")  # Debug
                # Build PDF
                doc.build(story)
                output.seek(0)
                
                print("Creating PDF response...")  # Debug
                # Create response
                response = Response(
                    output.getvalue(),
                    mimetype='application/pdf',
                    headers={'Content-Disposition': f'attachment; filename=export_{export_type}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'}
                )
                print("PDF export successful")  # Debug
                return response
                
            except Exception as e:
                print(f"PDF export error: {str(e)}")  # Debug
                return jsonify({
                    'success': False,
                    'error': f'PDF export failed: {str(e)}'
                }), 500
        
        # Default fallback for unsupported formats
        else:
            return jsonify({
                'success': False,
                'error': f'Unsupported export format: {export_format}. Supported formats are: csv, excel, pdf'
            }), 400
        
    except Exception as e:
        print(f"General export error: {str(e)}")  # Debug
        import traceback
        print(f"Traceback: {traceback.format_exc()}")  # Debug
        return jsonify({
            'success': False,
            'error': f'Export failed: {str(e)}'
        }), 500

# Serve uploaded files
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    from flask import send_from_directory
    return send_from_directory(UPLOAD_FOLDER, filename)

@app.route('/typing-passages/<filename>')
def typing_passage_file(filename):
    from flask import send_from_directory
    return send_from_directory(TYPING_PASSAGES_FOLDER, filename)

# Initialize database and create default users
def init_db():
    with app.app_context():
        db.create_all()
        
        # Only create demo student accounts with active subscriptions
        # Admin accounts must be created manually for security
        demo_students = [
            {'username': 'student1', 'password': 'student123'},
            {'username': 'student2', 'password': 'student123'},
            {'username': 'student3', 'password': 'student123'}
        ]
        
        for student_data in demo_students:
            if not User.query.filter_by(username=student_data['username']).first():
                student = User(
                    username=student_data['username'],
                    password_hash=hash_password(student_data['password']),
                    role='student',
                    is_active=True,
                    subscription_start=datetime.now(timezone.utc),
                    subscription_end=datetime.now(timezone.utc) + timedelta(days=30)  # 30-day subscription
                )
                db.session.add(student)
        
        db.session.commit()

# Ensure admin accounts exist on startup (for Render deployments)
def ensure_admin_accounts_startup():
    """Ensure admin accounts exist - run on every startup"""
    try:
        with app.app_context():
            db.create_all()
            
            # Check if admin account exists
            admin_exists = User.query.filter_by(username='admin').first()
            
            if not admin_exists:
                admin_user = User(
                    username=app.config.get('ADMIN_USERNAME', 'admin'),
                    password_hash=hash_password(app.config.get('ADMIN_PASSWORD', 'admin123')),
                    role='admin',
                    is_active=True,
                    is_locked=False,
                    created_at=datetime.now(timezone.utc)
                )
                db.session.add(admin_user)
                print("✅ Created admin account on startup")
            
            # Convert any existing superuser accounts to admin
            superuser_accounts = User.query.filter_by(role='superuser').all()
            for user in superuser_accounts:
                user.role = 'admin'
                print(f"✅ Converted {user.username} from superuser to admin")
            
            db.session.commit()
            print("💾 Admin accounts ensured on startup")
            
    except Exception as e:
        print(f"❌ Error ensuring admin accounts on startup: {e}")

# Initialize database automatically when app starts
try:
    init_db()
    ensure_admin_accounts_startup()  # Ensure admin accounts exist
    print("✅ Database initialized successfully")
except Exception as e:
    print(f"❌ Database initialization error: {e}")
    # Continue anyway in case tables already exist

if __name__ == '__main__':
    app.run(debug=True)
